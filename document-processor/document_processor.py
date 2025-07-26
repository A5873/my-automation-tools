#!/usr/bin/env python3
"""
Document Processor
A comprehensive toolkit for document manipulation, conversion, and automation.
"""

import os
import sys
import json
import argparse
import shutil
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional, Union
import tempfile

try:
    import pandas as pd
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    import PyPDF2
    import pdfplumber
    import docx2txt
    import markdown
    from bs4 import BeautifulSoup
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from fpdf import FPDF
    import openpyxl
    from jinja2 import Template
    import html2text
    import mammoth
    from tqdm import tqdm
    from colorama import init, Fore, Style
    import textstat
except ImportError as e:
    print(f"Missing required package: {e}")
    print("Install with: pip install -r requirements.txt")
    sys.exit(1)

# Initialize colorama
init(autoreset=True)

class DocumentProcessor:
    """Main class for processing and manipulating documents."""
    
    def __init__(self):
        self.supported_formats = {
            'input': ['.docx', '.pdf', '.txt', '.md', '.html', '.xlsx', '.csv'],
            'output': ['.docx', '.pdf', '.txt', '.md', '.html', '.xlsx', '.csv']
        }
        self.output_dir = Path("processed_documents")
        self.output_dir.mkdir(exist_ok=True)
        
        # Document templates
        self.templates_dir = Path("document_templates")
        self.templates_dir.mkdir(exist_ok=True)
        
        # Initialize document statistics
        self.stats = {
            'processed': 0,
            'converted': 0,
            'generated': 0,
            'errors': 0
        }
    
    def print_banner(self):
        """Display welcome banner."""
        print(f"\n{Fore.CYAN}{'='*60}")
        print(f"{Fore.CYAN}📄 DOCUMENT PROCESSOR SUITE 📄")
        print(f"{Fore.CYAN}{'='*60}")
        print(f"{Fore.YELLOW}Your comprehensive document automation toolkit!")
        print(f"{Fore.GREEN}Output directory: {self.output_dir.absolute()}")
        print(f"{Fore.CYAN}{'='*60}\n")
    
    def get_document_info(self, file_path: str) -> Dict:
        """Get comprehensive document information."""
        try:
            file_path = Path(file_path)
            info = {
                'filename': file_path.name,
                'extension': file_path.suffix.lower(),
                'size': file_path.stat().st_size,
                'modified': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
                'type': 'unknown'
            }
            
            # Determine document type and extract specific info
            if info['extension'] == '.pdf':
                info.update(self._get_pdf_info(file_path))
            elif info['extension'] == '.docx':
                info.update(self._get_docx_info(file_path))
            elif info['extension'] in ['.txt', '.md']:
                info.update(self._get_text_info(file_path))
            elif info['extension'] == '.xlsx':
                info.update(self._get_excel_info(file_path))
            
            return info
        except Exception as e:
            return {'error': str(e)}
    
    def _get_pdf_info(self, file_path: Path) -> Dict:
        """Get PDF-specific information."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return {
                    'type': 'PDF',
                    'pages': len(pdf_reader.pages),
                    'encrypted': pdf_reader.is_encrypted,
                    'metadata': dict(pdf_reader.metadata) if pdf_reader.metadata else {}
                }
        except Exception as e:
            return {'type': 'PDF', 'error': str(e)}
    
    def _get_docx_info(self, file_path: Path) -> Dict:
        """Get DOCX-specific information."""
        try:
            doc = Document(file_path)
            paragraphs = len(doc.paragraphs)
            tables = len(doc.tables)
            
            # Count words
            text = '\n'.join([p.text for p in doc.paragraphs])
            word_count = len(text.split())
            
            return {
                'type': 'Word Document',
                'paragraphs': paragraphs,
                'tables': tables,
                'word_count': word_count,
                'readability_score': textstat.flesch_reading_ease(text) if text.strip() else 0
            }
        except Exception as e:
            return {'type': 'Word Document', 'error': str(e)}
    
    def _get_text_info(self, file_path: Path) -> Dict:
        """Get text file information."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                lines = content.count('\n') + 1
                words = len(content.split())
                chars = len(content)
                
                return {
                    'type': 'Text File',
                    'lines': lines,
                    'words': words,
                    'characters': chars,
                    'readability_score': textstat.flesch_reading_ease(content) if content.strip() else 0
                }
        except Exception as e:
            return {'type': 'Text File', 'error': str(e)}
    
    def _get_excel_info(self, file_path: Path) -> Dict:
        """Get Excel file information."""
        try:
            workbook = openpyxl.load_workbook(file_path)
            sheets = len(workbook.sheetnames)
            
            # Count total rows and columns across all sheets
            total_rows = 0
            total_cols = 0
            
            for sheet in workbook.worksheets:
                total_rows += sheet.max_row
                total_cols = max(total_cols, sheet.max_column)
            
            return {
                'type': 'Excel Spreadsheet',
                'sheets': sheets,
                'sheet_names': workbook.sheetnames,
                'total_rows': total_rows,
                'max_columns': total_cols
            }
        except Exception as e:
            return {'type': 'Excel Spreadsheet', 'error': str(e)}
    
    def convert_document(self, input_path: str, output_format: str, output_path: str = None) -> str:
        """Convert document from one format to another."""
        input_path = Path(input_path)
        input_ext = input_path.suffix.lower()
        
        if not output_path:
            output_path = self.output_dir / f"{input_path.stem}_converted.{output_format}"
        else:
            output_path = Path(output_path)
        
        try:
            # PDF conversions
            if input_ext == '.pdf':
                if output_format == 'txt':
                    return self._pdf_to_text(input_path, output_path)
                elif output_format == 'docx':
                    return self._pdf_to_docx(input_path, output_path)
            
            # DOCX conversions
            elif input_ext == '.docx':
                if output_format == 'pdf':
                    return self._docx_to_pdf(input_path, output_path)
                elif output_format == 'txt':
                    return self._docx_to_text(input_path, output_path)
                elif output_format == 'html':
                    return self._docx_to_html(input_path, output_path)
                elif output_format == 'md':
                    return self._docx_to_markdown(input_path, output_path)
            
            # Text conversions
            elif input_ext == '.txt':
                if output_format == 'docx':
                    return self._text_to_docx(input_path, output_path)
                elif output_format == 'pdf':
                    return self._text_to_pdf(input_path, output_path)
                elif output_format == 'html':
                    return self._text_to_html(input_path, output_path)
            
            # Markdown conversions
            elif input_ext == '.md':
                if output_format == 'html':
                    return self._markdown_to_html(input_path, output_path)
                elif output_format == 'docx':
                    return self._markdown_to_docx(input_path, output_path)
                elif output_format == 'pdf':
                    return self._markdown_to_pdf(input_path, output_path)
            
            # HTML conversions
            elif input_ext == '.html':
                if output_format == 'txt':
                    return self._html_to_text(input_path, output_path)
                elif output_format == 'md':
                    return self._html_to_markdown(input_path, output_path)
                elif output_format == 'docx':
                    return self._html_to_docx(input_path, output_path)
            
            # Excel conversions
            elif input_ext == '.xlsx':
                if output_format == 'csv':
                    return self._excel_to_csv(input_path, output_path)
                elif output_format == 'txt':
                    return self._excel_to_text(input_path, output_path)
            
            else:
                raise ValueError(f"Conversion from {input_ext} to {output_format} not supported")
                
        except Exception as e:
            self.stats['errors'] += 1
            raise Exception(f"Conversion failed: {str(e)}")
    
    def _pdf_to_text(self, input_path: Path, output_path: Path) -> str:
        """Convert PDF to text."""
        with pdfplumber.open(input_path) as pdf:
            text_content = []
            for page in pdf.pages:
                text_content.append(page.extract_text() or "")
        
        full_text = '\n\n--- Page Break ---\n\n'.join(text_content)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_text)
        
        self.stats['converted'] += 1
        return str(output_path)
    
    def _docx_to_text(self, input_path: Path, output_path: Path) -> str:
        """Convert DOCX to text."""
        text = docx2txt.process(str(input_path))
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        self.stats['converted'] += 1
        return str(output_path)
    
    def _docx_to_html(self, input_path: Path, output_path: Path) -> str:
        """Convert DOCX to HTML using mammoth."""
        with open(input_path, 'rb') as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html_content = result.value
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        self.stats['converted'] += 1
        return str(output_path)
    
    def _text_to_docx(self, input_path: Path, output_path: Path) -> str:
        """Convert text to DOCX."""
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Document: {input_path.name}', 0)
        
        # Add content paragraphs
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        doc.save(output_path)
        self.stats['converted'] += 1
        return str(output_path)
    
    def _text_to_pdf(self, input_path: Path, output_path: Path) -> str:
        """Convert text to PDF."""
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        # Split content into lines that fit the page
        lines = content.split('\n')
        for line in lines:
            # Handle long lines by wrapping them
            if len(line) > 80:
                words = line.split(' ')
                current_line = ""
                for word in words:
                    if len(current_line + word + " ") <= 80:
                        current_line += word + " "
                    else:
                        if current_line:
                            pdf.cell(0, 10, current_line.strip(), ln=1)
                        current_line = word + " "
                if current_line:
                    pdf.cell(0, 10, current_line.strip(), ln=1)
            else:
                pdf.cell(0, 10, line, ln=1)
        
        pdf.output(str(output_path))
        self.stats['converted'] += 1
        return str(output_path)
    
    def _markdown_to_html(self, input_path: Path, output_path: Path) -> str:
        """Convert Markdown to HTML."""
        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        html_content = markdown.markdown(md_content, extensions=['tables', 'codehilite'])
        
        # Wrap in basic HTML structure
        full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{input_path.stem}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; }}
        code {{ background-color: #f4f4f4; padding: 2px 4px; }}
        pre {{ background-color: #f4f4f4; padding: 10px; overflow-x: auto; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_html)
        
        self.stats['converted'] += 1
        return str(output_path)
    
    def _html_to_text(self, input_path: Path, output_path: Path) -> str:
        """Convert HTML to text."""
        with open(input_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Convert HTML to text
        h = html2text.HTML2Text()
        h.ignore_links = False
        text_content = h.handle(html_content)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
        
        self.stats['converted'] += 1
        return str(output_path)
    
    def _excel_to_csv(self, input_path: Path, output_path: Path) -> str:
        """Convert Excel to CSV."""
        # Read all sheets and combine or save separately
        excel_file = pd.ExcelFile(input_path)
        
        if len(excel_file.sheet_names) == 1:
            # Single sheet - save as single CSV
            df = pd.read_excel(input_path)
            df.to_csv(output_path, index=False)
        else:
            # Multiple sheets - save each as separate CSV
            output_dir = output_path.parent / f"{output_path.stem}_sheets"
            output_dir.mkdir(exist_ok=True)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(input_path, sheet_name=sheet_name)
                sheet_csv = output_dir / f"{sheet_name}.csv"
                df.to_csv(sheet_csv, index=False)
        
        self.stats['converted'] += 1
        return str(output_path)
    
    def _excel_to_text(self, input_path: Path, output_path: Path) -> str:
        """Convert Excel to text."""
        excel_file = pd.ExcelFile(input_path)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            for i, sheet_name in enumerate(excel_file.sheet_names):
                if i > 0:
                    f.write('\n\n' + '='*50 + '\n\n')
                
                f.write(f"Sheet: {sheet_name}\n")
                f.write('-' * 30 + '\n\n')
                
                df = pd.read_excel(input_path, sheet_name=sheet_name)
                f.write(df.to_string(index=False))
                f.write('\n')
        
        self.stats['converted'] += 1
        return str(output_path)
    
    def _pdf_to_docx(self, input_path: Path, output_path: Path) -> str:
        """Convert PDF to DOCX."""
        # Extract text from PDF
        text_content = []
        with pdfplumber.open(input_path) as pdf:
            for page in pdf.pages:
                text_content.append(page.extract_text() or "")
        
        # Create DOCX document
        doc = Document()
        doc.add_heading(f'Converted from: {input_path.name}', 0)
        
        for i, page_text in enumerate(text_content):
            if i > 0:
                doc.add_page_break()
            
            doc.add_heading(f'Page {i+1}', 1)
            
            # Add text paragraphs
            paragraphs = page_text.split('\n\n') if page_text else []
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
        
        doc.save(output_path)
        self.stats['converted'] += 1
        return str(output_path)
    
    def _docx_to_pdf(self, input_path: Path, output_path: Path) -> str:
        """Convert DOCX to PDF using reportlab."""
        # Extract text from DOCX
        text = docx2txt.process(str(input_path))
        
        # Create PDF
        doc = SimpleDocTemplate(str(output_path), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        title = Paragraph(f"Converted from: {input_path.name}", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 20))
        
        # Add content paragraphs
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = Paragraph(para.strip().replace('\n', '<br/>'), styles['Normal'])
                story.append(p)
                story.append(Spacer(1, 12))
        
        doc.build(story)
        self.stats['converted'] += 1
        return str(output_path)
    
    def _docx_to_markdown(self, input_path: Path, output_path: Path) -> str:
        """Convert DOCX to Markdown."""
        # Extract text from DOCX
        doc = Document(input_path)
        
        markdown_content = []
        markdown_content.append(f"# {input_path.stem}\n")
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Check if it's a heading (simple heuristic)
                if paragraph.style.name.startswith('Heading'):
                    level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
                    markdown_content.append('#' * level + ' ' + text + '\n')
                else:
                    markdown_content.append(text + '\n')
        
        # Handle tables
        for table in doc.tables:
            markdown_content.append('\n')
            for i, row in enumerate(table.rows):
                row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                markdown_content.append('| ' + row_text + ' |\n')
                
                # Add header separator for first row
                if i == 0:
                    separator = '|' + '---|' * len(row.cells)
                    markdown_content.append(separator + '\n')
            markdown_content.append('\n')
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(''.join(markdown_content))
        
        self.stats['converted'] += 1
        return str(output_path)
    
    def _markdown_to_docx(self, input_path: Path, output_path: Path) -> str:
        """Convert Markdown to DOCX."""
        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert markdown to HTML first
        html_content = markdown.markdown(md_content, extensions=['tables'])
        
        # Parse HTML and create DOCX
        soup = BeautifulSoup(html_content, 'html.parser')
        doc = Document()
        
        doc.add_heading(f'Converted from: {input_path.name}', 0)
        
        # Process HTML elements
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol']):
            if element.name.startswith('h'):
                level = int(element.name[1])
                doc.add_heading(element.get_text().strip(), level)
            elif element.name == 'p':
                text = element.get_text().strip()
                if text:
                    doc.add_paragraph(text)
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    doc.add_paragraph(li.get_text().strip(), style='List Bullet')
        
        doc.save(output_path)
        self.stats['converted'] += 1
        return str(output_path)
    
    def _markdown_to_pdf(self, input_path: Path, output_path: Path) -> str:
        """Convert Markdown to PDF."""
        # First convert to HTML
        temp_html = tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False)
        self._markdown_to_html(input_path, Path(temp_html.name))
        
        try:
            # Then convert HTML to text and create PDF
            with open(temp_html.name, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # Extract text from HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            text_content = soup.get_text()
            
            # Create PDF using fpdf
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            # Add content
            lines = text_content.split('\n')
            for line in lines:
                if line.strip():
                    # Handle long lines
                    if len(line) > 80:
                        words = line.split(' ')
                        current_line = ""
                        for word in words:
                            if len(current_line + word + " ") <= 80:
                                current_line += word + " "
                            else:
                                if current_line:
                                    pdf.cell(0, 10, current_line.strip(), ln=1)
                                current_line = word + " "
                        if current_line:
                            pdf.cell(0, 10, current_line.strip(), ln=1)
                    else:
                        pdf.cell(0, 10, line, ln=1)
            
            pdf.output(str(output_path))
            
        finally:
            # Clean up temporary file
            os.unlink(temp_html.name)
        
        self.stats['converted'] += 1
        return str(output_path)
    
    def _html_to_markdown(self, input_path: Path, output_path: Path) -> str:
        """Convert HTML to Markdown."""
        with open(input_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Convert HTML to Markdown
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.body_width = 0  # No line wrapping
        markdown_content = h.handle(html_content)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        self.stats['converted'] += 1
        return str(output_path)
    
    def _html_to_docx(self, input_path: Path, output_path: Path) -> str:
        """Convert HTML to DOCX."""
        with open(input_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Parse HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        doc = Document()
        
        doc.add_heading(f'Converted from: {input_path.name}', 0)
        
        # Process HTML elements
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'div']):
            if element.name.startswith('h'):
                level = int(element.name[1])
                text = element.get_text().strip()
                if text:
                    doc.add_heading(text, level)
            elif element.name == 'p':
                text = element.get_text().strip()
                if text:
                    doc.add_paragraph(text)
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    text = li.get_text().strip()
                    if text:
                        doc.add_paragraph(text, style='List Bullet')
            elif element.name == 'div':
                text = element.get_text().strip()
                if text and not element.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol']):
                    doc.add_paragraph(text)
        
        doc.save(output_path)
        self.stats['converted'] += 1
        return str(output_path)
    
    def _text_to_html(self, input_path: Path, output_path: Path) -> str:
        """Convert text to HTML."""
        with open(input_path, 'r', encoding='utf-8') as f:
            text_content = f.read()
        
        # Create HTML
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{input_path.stem}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        pre {{ background-color: #f4f4f4; padding: 10px; overflow-x: auto; }}
    </style>
</head>
<body>
    <h1>{input_path.stem}</h1>
    <pre>{text_content}</pre>
</body>
</html>"""
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        self.stats['converted'] += 1
        return str(output_path)
    
    def generate_document_from_template(self, template_path: str, data_dict: Dict, output_path: str) -> str:
        """Generate document from template using Jinja2."""
        template_path = Path(template_path)
        
        with open(template_path, 'r', encoding='utf-8') as f:
            template_content = f.read()
        
        template = Template(template_content)
        rendered_content = template.render(**data_dict)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(rendered_content)
        
        self.stats['generated'] += 1
        return str(output_path)
    
    def merge_documents(self, file_paths: List[str], output_path: str, format_type: str = 'docx') -> str:
        """Merge multiple documents into one."""
        if format_type == 'docx':
            return self._merge_docx_documents(file_paths, output_path)
        elif format_type == 'pdf':
            return self._merge_pdf_documents(file_paths, output_path)
        elif format_type == 'txt':
            return self._merge_text_documents(file_paths, output_path)
        else:
            raise ValueError(f"Merging format {format_type} not supported")
    
    def _merge_docx_documents(self, file_paths: List[str], output_path: str) -> str:
        """Merge DOCX documents."""
        merged_doc = Document()
        
        for i, file_path in enumerate(file_paths):
            if i > 0:
                merged_doc.add_page_break()
            
            # Add document title
            merged_doc.add_heading(f'Document {i+1}: {Path(file_path).name}', level=1)
            
            # Read and add content
            source_doc = Document(file_path)
            for paragraph in source_doc.paragraphs:
                new_para = merged_doc.add_paragraph()
                new_para.text = paragraph.text
                # Copy formatting if possible
                new_para.style = paragraph.style
        
        merged_doc.save(output_path)
        self.stats['processed'] += 1
        return str(output_path)
    
    def _merge_pdf_documents(self, file_paths: List[str], output_path: str) -> str:
        """Merge PDF documents."""
        pdf_writer = PyPDF2.PdfWriter()
        
        for file_path in file_paths:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    pdf_writer.add_page(page)
        
        with open(output_path, 'wb') as output_file:
            pdf_writer.write(output_file)
        
        self.stats['processed'] += 1
        return str(output_path)
    
    def _merge_text_documents(self, file_paths: List[str], output_path: str) -> str:
        """Merge text documents."""
        with open(output_path, 'w', encoding='utf-8') as output_file:
            for i, file_path in enumerate(file_paths):
                if i > 0:
                    output_file.write('\n\n' + '='*50 + '\n\n')
                
                output_file.write(f"Document {i+1}: {Path(file_path).name}\n")
                output_file.write('-' * 30 + '\n\n')
                
                with open(file_path, 'r', encoding='utf-8') as input_file:
                    content = input_file.read()
                    output_file.write(content)
        
        self.stats['processed'] += 1
        return str(output_path)
    
    def extract_text_from_document(self, file_path: str) -> str:
        """Extract text content from various document formats."""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                with pdfplumber.open(file_path) as pdf:
                    text_content = []
                    for page in pdf.pages:
                        text_content.append(page.extract_text() or "")
                    return '\n\n'.join(text_content)
            
            elif extension == '.docx':
                return docx2txt.process(str(file_path))
            
            elif extension in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            elif extension == '.html':
                with open(file_path, 'r', encoding='utf-8') as f:
                    soup = BeautifulSoup(f.read(), 'html.parser')
                    return soup.get_text()
            
            else:
                raise ValueError(f"Text extraction from {extension} not supported")
                
        except Exception as e:
            raise Exception(f"Text extraction failed: {str(e)}")
    
    def analyze_document(self, file_path: str) -> Dict:
        """Analyze document and provide detailed statistics."""
        try:
            text_content = self.extract_text_from_document(file_path)
            
            # Basic statistics
            word_count = len(text_content.split())
            char_count = len(text_content)
            char_count_no_spaces = len(text_content.replace(' ', ''))
            paragraph_count = len([p for p in text_content.split('\n\n') if p.strip()])
            sentence_count = text_content.count('.') + text_content.count('!') + text_content.count('?')
            
            # Readability statistics
            analysis = {
                'file_info': self.get_document_info(file_path),
                'word_count': word_count,
                'character_count': char_count,
                'character_count_no_spaces': char_count_no_spaces,
                'paragraph_count': paragraph_count,
                'sentence_count': sentence_count,
                'average_words_per_sentence': word_count / sentence_count if sentence_count > 0 else 0,
                'readability': {
                    'flesch_reading_ease': textstat.flesch_reading_ease(text_content),
                    'flesch_kincaid_grade': textstat.flesch_kincaid_grade(text_content),
                    'automated_readability_index': textstat.automated_readability_index(text_content),
                    'reading_time_minutes': word_count / 200  # Average reading speed
                }
            }
            
            return analysis
            
        except Exception as e:
            return {'error': str(e)}
    
    def batch_convert_documents(self, input_dir: str, output_format: str, filter_extensions: List[str] = None) -> Dict:
        """Batch convert all documents in a directory."""
        input_dir = Path(input_dir)
        results = {
            'successful': [],
            'failed': [],
            'total_processed': 0
        }
        
        # Find all supported documents
        if filter_extensions:
            document_files = []
            for ext in filter_extensions:
                document_files.extend(input_dir.glob(f"*{ext}"))
        else:
            document_files = []
            for ext in self.supported_formats['input']:
                document_files.extend(input_dir.glob(f"*{ext}"))
        
        if not document_files:
            return results
        
        print(f"{Fore.CYAN}🔄 Processing {len(document_files)} documents...")
        
        for doc_file in tqdm(document_files, desc="Converting documents"):
            try:
                output_path = self.output_dir / f"{doc_file.stem}_converted.{output_format}"
                converted_path = self.convert_document(str(doc_file), output_format, str(output_path))
                results['successful'].append({
                    'input': str(doc_file),
                    'output': converted_path
                })
            except Exception as e:
                results['failed'].append({
                    'input': str(doc_file),
                    'error': str(e)
                })
            
            results['total_processed'] += 1
        
        return results
    
    def create_document_report(self, file_paths: List[str], output_path: str) -> str:
        """Create a comprehensive report analyzing multiple documents."""
        report_data = {
            'generated_at': datetime.now().isoformat(),
            'total_documents': len(file_paths),
            'documents': []
        }
        
        print(f"{Fore.CYAN}📊 Analyzing {len(file_paths)} documents...")
        
        for file_path in tqdm(file_paths, desc="Analyzing documents"):
            try:
                analysis = self.analyze_document(file_path)
                report_data['documents'].append(analysis)
            except Exception as e:
                report_data['documents'].append({
                    'file_path': file_path,
                    'error': str(e)
                })
        
        # Generate HTML report
        html_template = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>Document Analysis Report</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 40px; }
                .header { background-color: #f0f0f0; padding: 20px; border-radius: 5px; }
                .document { margin: 20px 0; padding: 15px; border: 1px solid #ddd; border-radius: 5px; }
                .stats { display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 10px; }
                .stat-item { background-color: #f9f9f9; padding: 10px; border-radius: 3px; }
                .error { background-color: #ffebee; color: #c62828; }
            </style>
        </head>
        <body>
            <div class="header">
                <h1>📄 Document Analysis Report</h1>
                <p>Generated: {{ generated_at }}</p>
                <p>Total Documents: {{ total_documents }}</p>
            </div>
            
            {% for doc in documents %}
            <div class="document">
                {% if doc.error %}
                    <div class="error">
                        <h3>❌ Error processing document</h3>
                        <p>{{ doc.error }}</p>
                    </div>
                {% else %}
                    <h3>📄 {{ doc.file_info.filename }}</h3>
                    <div class="stats">
                        <div class="stat-item">
                            <strong>Word Count:</strong> {{ doc.word_count }}
                        </div>
                        <div class="stat-item">
                            <strong>Characters:</strong> {{ doc.character_count }}
                        </div>
                        <div class="stat-item">
                            <strong>Paragraphs:</strong> {{ doc.paragraph_count }}
                        </div>
                        <div class="stat-item">
                            <strong>Reading Time:</strong> {{ "%.1f"|format(doc.readability.reading_time_minutes) }} min
                        </div>
                        <div class="stat-item">
                            <strong>Reading Level:</strong> Grade {{ "%.1f"|format(doc.readability.flesch_kincaid_grade) }}
                        </div>
                        <div class="stat-item">
                            <strong>File Size:</strong> {{ "%.2f"|format(doc.file_info.size / 1024) }} KB
                        </div>
                    </div>
                {% endif %}
            </div>
            {% endfor %}
        </body>
        </html>
        """
        
        template = Template(html_template)
        html_content = template.render(**report_data)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        # Also save JSON data
        json_path = Path(output_path).with_suffix('.json')
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(report_data, f, indent=2, ensure_ascii=False)
        
        self.stats['generated'] += 1
        return str(output_path)

def main():
    """Main function for CLI usage."""
    parser = argparse.ArgumentParser(description='Document Processor')
    parser.add_argument('--mode', choices=['convert', 'merge', 'analyze', 'batch', 'report'], 
                       default='convert', help='Operation mode')
    parser.add_argument('--input', help='Input file or directory path')
    parser.add_argument('--output', help='Output file path')
    parser.add_argument('--format', help='Output format')
    parser.add_argument('--files', nargs='+', help='Multiple input files for merging')
    
    args = parser.parse_args()
    
    processor = DocumentProcessor()
    processor.print_banner()
    
    try:
        if args.mode == 'convert' and args.input and args.format:
            result = processor.convert_document(args.input, args.format, args.output)
            print(f"{Fore.GREEN}✅ Document converted: {result}")
        
        elif args.mode == 'merge' and args.files and args.output:
            result = processor.merge_documents(args.files, args.output, args.format or 'docx')
            print(f"{Fore.GREEN}✅ Documents merged: {result}")
        
        elif args.mode == 'analyze' and args.input:
            analysis = processor.analyze_document(args.input)
            print(f"{Fore.CYAN}📊 Document Analysis:")
            print(json.dumps(analysis, indent=2))
        
        elif args.mode == 'batch' and args.input and args.format:
            results = processor.batch_convert_documents(args.input, args.format)
            print(f"{Fore.GREEN}✅ Batch conversion completed:")
            print(f"   Successful: {len(results['successful'])}")
            print(f"   Failed: {len(results['failed'])}")
        
        elif args.mode == 'report' and args.files:
            output_path = args.output or 'document_report.html'
            result = processor.create_document_report(args.files, output_path)
            print(f"{Fore.GREEN}✅ Report generated: {result}")
        
        else:
            parser.print_help()
            
    except Exception as e:
        print(f"{Fore.RED}❌ Error: {e}")

if __name__ == "__main__":
    main()
